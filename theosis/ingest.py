"""Lớp tri giác (perception) — biến mọi định dạng file thành TEXT dùng chung.

Triết lý: council gồm các model dị chủng (có model thấy ảnh, có model mù; PDF thì
gần như không model nào đọc raw). Để giữ bất biến MoA — *mọi thành viên nghị bàn
trên cùng một đầu vào* — ta chuẩn hoá mọi file về text ở một lớp tiền xử lý, rồi
chia khối text đó cho tất cả slot.

Thuần & đồng bộ cho các định dạng file (không gọi model). Ảnh trả về metadata +
cờ needs_vision; phần mô tả ảnh bằng model thị giác là bước async riêng ở server.

Thư viện đọc tài liệu là tuỳ chọn (extra `[files]`): pypdf, python-docx, openpyxl,
Pillow. Thiếu lib nào thì định dạng đó trả về note nhẹ nhàng thay vì lỗi.
"""
from __future__ import annotations

import csv
import io
import json
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import List

MAX_CHARS = 20_000          # cắt mỗi file để khỏi phình token
MAX_TABLE_ROWS = 200

TEXT_EXT = {".txt", ".md", ".markdown", ".rst", ".log"}
CODE_EXT = {".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".cpp", ".h", ".hpp",
            ".go", ".rs", ".rb", ".php", ".sh", ".sql", ".html", ".css", ".xml",
            ".toml", ".ini", ".cfg"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}


@dataclass
class IngestResult:
    filename: str
    kind: str                       # pdf|docx|xlsx|csv|text|code|json|yaml|image|unknown
    text: str = ""
    chars: int = 0
    truncated: bool = False
    note: str = ""
    needs_vision: bool = False
    meta: dict = field(default_factory=dict)

    def as_dict(self) -> dict:
        return asdict(self)


def _truncate(text: str) -> tuple[str, bool]:
    if len(text) > MAX_CHARS:
        return text[:MAX_CHARS] + "\n…[đã cắt bớt]", True
    return text, False


def _ext(filename: str) -> str:
    return Path(filename).suffix.lower()


# ── từng định dạng (đều nhận bytes) ──────────────────────────────────────────
def _read_text_bytes(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="replace")


def _ingest_pdf(data: bytes, filename: str) -> IngestResult:
    try:
        from pypdf import PdfReader
    except Exception:
        return IngestResult(filename, "pdf", note="Cần `pip install pypdf` để đọc PDF.")
    try:
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for i, page in enumerate(reader.pages):
            t = (page.extract_text() or "").strip()
            if t:
                parts.append(f"--- trang {i + 1} ---\n{t}")
        text, trunc = _truncate("\n\n".join(parts))
        note = "" if parts else "PDF không trích được text (có thể là scan ảnh — cần OCR)."
        return IngestResult(filename, "pdf", text=text, chars=len(text), truncated=trunc,
                            note=note, meta={"pages": len(reader.pages)})
    except Exception as e:
        return IngestResult(filename, "pdf", note=f"Lỗi đọc PDF: {e}")


def _ingest_docx(data: bytes, filename: str) -> IngestResult:
    try:
        import docx
    except Exception:
        return IngestResult(filename, "docx", note="Cần `pip install python-docx` để đọc .docx.")
    try:
        doc = docx.Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        text, trunc = _truncate("\n".join(lines))
        return IngestResult(filename, "docx", text=text, chars=len(text), truncated=trunc)
    except Exception as e:
        return IngestResult(filename, "docx", note=f"Lỗi đọc .docx: {e}")


def _ingest_xlsx(data: bytes, filename: str) -> IngestResult:
    try:
        import openpyxl
    except Exception:
        return IngestResult(filename, "xlsx", note="Cần `pip install openpyxl` để đọc .xlsx.")
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        blocks = []
        for ws in wb.worksheets:
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= MAX_TABLE_ROWS:
                    rows.append("…[còn dòng]")
                    break
                rows.append(" | ".join("" if c is None else str(c) for c in row))
            blocks.append(f"### Sheet: {ws.title}\n" + "\n".join(rows))
        wb.close()
        text, trunc = _truncate("\n\n".join(blocks))
        return IngestResult(filename, "xlsx", text=text, chars=len(text), truncated=trunc)
    except Exception as e:
        return IngestResult(filename, "xlsx", note=f"Lỗi đọc .xlsx: {e}")


def _ingest_csv(data: bytes, filename: str, delim: str = ",") -> IngestResult:
    try:
        raw = _read_text_bytes(data)
        reader = csv.reader(io.StringIO(raw), delimiter=delim)
        rows = []
        for i, row in enumerate(reader):
            if i >= MAX_TABLE_ROWS:
                rows.append("…[còn dòng]")
                break
            rows.append(" | ".join(row))
        text, trunc = _truncate("\n".join(rows))
        return IngestResult(filename, "csv", text=text, chars=len(text), truncated=trunc,
                            meta={"rows": len(rows)})
    except Exception as e:
        return IngestResult(filename, "csv", note=f"Lỗi đọc CSV: {e}")


def _ingest_json(data: bytes, filename: str) -> IngestResult:
    raw = _read_text_bytes(data)
    try:
        obj = json.loads(raw)
        pretty = json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        pretty = raw  # không phải JSON hợp lệ → giữ nguyên text
    text, trunc = _truncate(pretty)
    return IngestResult(filename, "json", text=text, chars=len(text), truncated=trunc)


def _ingest_plain(data: bytes, filename: str, kind: str) -> IngestResult:
    text, trunc = _truncate(_read_text_bytes(data))
    return IngestResult(filename, kind, text=text, chars=len(text), truncated=trunc)


def _ingest_image(data: bytes, filename: str) -> IngestResult:
    w = h = None
    try:
        from PIL import Image
        with Image.open(io.BytesIO(data)) as im:
            w, h = im.size
    except Exception:
        pass
    dim = f"{w}×{h}" if w else "kích thước không rõ"
    return IngestResult(
        filename, "image",
        text=f"[ảnh {filename} · {dim} · {len(data)} bytes — nội dung cần model thị giác mô tả]",
        note="Ảnh: cần một slot có thị giác để mô tả/OCR thành text.",
        needs_vision=True, meta={"width": w, "height": h, "bytes": len(data)},
    )


# ── điểm vào ─────────────────────────────────────────────────────────────────
def ingest_bytes(data: bytes, filename: str) -> IngestResult:
    """Biến (bytes, tên file) thành IngestResult chứa text dùng chung. Thuần, đồng bộ."""
    ext = _ext(filename)
    if ext == ".pdf":
        return _ingest_pdf(data, filename)
    if ext == ".docx":
        return _ingest_docx(data, filename)
    if ext in (".xlsx", ".xlsm"):
        return _ingest_xlsx(data, filename)
    if ext == ".csv":
        return _ingest_csv(data, filename, ",")
    if ext == ".tsv":
        return _ingest_csv(data, filename, "\t")
    if ext == ".json":
        return _ingest_json(data, filename)
    if ext in (".yaml", ".yml"):
        return _ingest_plain(data, filename, "yaml")
    if ext in IMAGE_EXT:
        return _ingest_image(data, filename)
    if ext in TEXT_EXT:
        return _ingest_plain(data, filename, "text")
    if ext in CODE_EXT:
        return _ingest_plain(data, filename, "code")
    # fallback: thử đọc như text; nếu nhị phân thì báo note
    try:
        raw = data.decode("utf-8")
        text, trunc = _truncate(raw)
        return IngestResult(filename, "text", text=text, chars=len(text), truncated=trunc,
                            note="Định dạng không rõ — đọc như văn bản thuần.")
    except Exception:
        return IngestResult(filename, "unknown", note=f"Không đọc được định dạng này ({ext or 'không đuôi'}).")


def ingest_path(path: str) -> IngestResult:
    p = Path(path)
    return ingest_bytes(p.read_bytes(), p.name)


def format_attachments(results: List[IngestResult]) -> str:
    """Ghép các IngestResult thành một khối context chia chung cho mọi slot."""
    if not results:
        return ""
    blocks = ["[TÀI LIỆU ĐÍNH KÈM]"]
    for r in results:
        head = f"=== {r.filename} ({r.kind}" + (f", {r.chars} ký tự" if r.chars else "") + ") ==="
        body = r.text if r.text else (r.note or "(không có nội dung)")
        blocks.append(head + "\n" + body)
    blocks.append("[HẾT TÀI LIỆU]")
    return "\n\n".join(blocks)


# ── trích artifact RA file (code block trong câu trả lời cuối) ────────────────
_LANG_EXT = {
    "python": ".py", "py": ".py", "javascript": ".js", "js": ".js", "typescript": ".ts",
    "ts": ".ts", "java": ".java", "c": ".c", "cpp": ".cpp", "csharp": ".cs", "go": ".go",
    "rust": ".rs", "ruby": ".rb", "php": ".php", "bash": ".sh", "sh": ".sh", "shell": ".sh",
    "sql": ".sql", "html": ".html", "css": ".css", "json": ".json", "yaml": ".yaml",
    "yml": ".yaml", "xml": ".xml", "markdown": ".md", "md": ".md", "csv": ".csv", "": ".txt",
}


def extract_code_blocks(text: str) -> List[dict]:
    """Tìm các code block ```lang ... ``` trong text → danh sách {lang, ext, content}.

    Dùng để xuất artifact ra file tải về. Thuần (không regex tham lam — quét theo dòng)."""
    out: List[dict] = []
    lines = (text or "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            content = "\n".join(buf).strip("\n")
            if content:
                ext = _LANG_EXT.get(lang, ".txt")
                out.append({"lang": lang or "text", "ext": ext, "content": content})
            i += 1  # bỏ qua ``` đóng
        else:
            i += 1
    return out
